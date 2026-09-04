"""
"Word 다운로드" 버튼 — /api/agent/build-candidate-docx, renderers.docx_renderer.

renderers/candidate_specification.py의 공통 Structured Data를 render_candidate_
markdown()과 함께 소비하므로(tests/test_specification_consistency.py가 둘의
일치를 검증한다), 이 파일은 Word 쪽에 특화된 것만 확인한다:

1. render_candidate_docx()가 실제로 열 수 있는(빈 파일이 아닌) .docx 바이트를
   만드는지 — python-docx로 다시 읽어 내용을 확인한다.
2. 장비명/주요 섹션/Requirement Compliance(PASS/FAIL/UNKNOWN)이 문서에 포함되는지.
3. /api/agent/build-candidate-docx가 실제로 파일을 만들고 다운로드까지 되는지.
4. 파일명이 장비명 기반으로 만들어지고, 이름이 없으면 fallback을 쓰는지.
"""
from __future__ import annotations

import io

import pytest
from docx import Document
from fastapi.testclient import TestClient

import main
from agent.schemas import CandidateEquipment, CandidateEquipmentFact, ComplianceRecord, RequirementSchema, SourceRef
from renderers.docx_renderer import render_candidate_docx

pytestmark = [pytest.mark.specification, pytest.mark.download]


def _client():
    return TestClient(main.app)


def _full_candidate() -> CandidateEquipment:
    return CandidateEquipment(
        candidate_id="cand-1",
        manufacturer="MultiSense",
        model="MS-600",
        source_document="SPEC-010.md",
        equipment_fact=CandidateEquipmentFact(
            equipment_type="Multi-Modal Electrode Inspection System",
            measurement_principle="3D Laser + Vision",
            inline_offline="inline",
            measurement_method="non_contact",
            width_mm=800.0,
            range_min=0.0,
            range_max=300.0,
            range_unit="um",
            accuracy_value=0.8,
            accuracy_unit="um",
            resolution_value=0.2,
            resolution_unit="um",
            speed_value=500.0,
            speed_unit="mm/s",
            defect_types=["Scratch", "Crack", "Particle", "Coating Defect"],
            min_defect_size_value=15.0,
            min_defect_size_unit="um",
        ),
    )


def _hard_requirement_report() -> list[ComplianceRecord]:
    return [
        ComplianceRecord(
            item="Width", unit="mm", requirement=600.0, specification=800.0, operator=">=",
            result="PASS", reason="장비 대응 폭 800mm >= 요구 600mm → PASS",
            source=SourceRef(document="SPEC-010.md"), hard=True,
        ),
        ComplianceRecord(
            item="Accuracy", unit="um", requirement=1.0, specification=0.8, operator="<=",
            result="PASS", reason="장비 정확도 0.8um <= 요구 1.0um → PASS",
            source=SourceRef(document="SPEC-010.md"), hard=True,
        ),
        ComplianceRecord(
            item="Repeatability", unit="um", requirement=0.5, specification=None, operator="<=",
            result="UNKNOWN", reason="사양서에서 확인할 수 없음", hard=True,
        ),
        ComplianceRecord(
            item="Speed", unit="mm/s", requirement=600.0, specification=500.0, operator=">=",
            result="FAIL", reason="장비 속도 500mm/s < 요구 600mm/s → FAIL",
            source=SourceRef(document="SPEC-010.md"), hard=True,
        ),
    ]


def _open_docx(docx_bytes: bytes) -> Document:
    assert len(docx_bytes) > 0, "생성된 .docx가 빈 파일임"
    return Document(io.BytesIO(docx_bytes))


def test_render_candidate_docx_produces_a_real_openable_document():
    candidate = _full_candidate()
    requirement = RequirementSchema(inspection_items=["thickness", "surface_defect"])
    docx_bytes = render_candidate_docx(candidate, requirement=requirement, hard_requirement_report=_hard_requirement_report())

    document = _open_docx(docx_bytes)
    assert len(document.paragraphs) > 0
    assert len(document.tables) > 0


def test_render_candidate_docx_includes_equipment_name_and_sections():
    candidate = _full_candidate()
    docx_bytes = render_candidate_docx(candidate, requirement=None, hard_requirement_report=None)
    document = _open_docx(docx_bytes)

    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "MultiSense MS-600" in full_text  # Title(장비명)

    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    for expected in (
        "General Specification", "Inspection Target", "Inspection Requirements",
        "Measurement Performance", "Spatial Performance", "Optical System",
        "Defect Inspection", "Inspection Performance", "System Configuration",
        "Requirement Compliance", "Sources / Notes",
    ):
        assert expected in headings, f"'{expected}' 섹션(Heading)이 문서에 없음: {headings}"


def test_render_candidate_docx_excludes_interfaces_environment_safety_sections():
    """요청서: 전극 검사 핵심 비교와 무관하고(CandidateEquipmentFact가 애초에
    추출하지 않아 항상 UNKNOWN뿐인) Interfaces/Data, Environment, Safety
    섹션은 Word 사양서에서 제외한다. Markdown(render_candidate_markdown)은
    이 변경의 영향을 받지 않고 기존처럼 13개 섹션을 그대로 유지한다
    (tests/test_specification_consistency.py 및 아래 markdown 쪽 회귀 테스트로
    별도 확인)."""
    candidate = _full_candidate()
    docx_bytes = render_candidate_docx(candidate, requirement=None, hard_requirement_report=None)
    document = _open_docx(docx_bytes)

    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    for excluded in ("Interfaces / Data", "Environment", "Safety"):
        assert excluded not in headings, f"'{excluded}' 섹션이 Word 사양서에서 제외되지 않음: {headings}"


def test_render_candidate_docx_requirement_compliance_shows_pass_fail_unknown_distinctly():
    candidate = _full_candidate()
    docx_bytes = render_candidate_docx(candidate, requirement=None, hard_requirement_report=_hard_requirement_report())
    document = _open_docx(docx_bytes)

    compliance_table = None
    for table in document.tables:
        header = [c.text for c in table.rows[0].cells]
        if header == ["Requirement", "Required", "Equipment", "Result"]:
            compliance_table = table
            break
    assert compliance_table is not None, "Requirement Compliance 표를 찾지 못함"

    rows = {row.cells[0].text: [c.text for c in row.cells] for row in compliance_table.rows[1:]}
    assert rows["Width"][3] == "PASS"
    assert rows["Accuracy"][3] == "PASS"
    assert rows["Repeatability"][3] == "UNKNOWN"
    assert rows["Speed"][3] == "FAIL"
    # UNKNOWN이 PASS처럼 보이면 안 된다 — Equipment 컬럼도 명확히 UNKNOWN이어야 함.
    assert rows["Repeatability"][2] == "UNKNOWN"
    # 값과 상태가 서로 다른 셀에 들어가 "0~300umVERIFIED"처럼 붙어보이지 않아야 한다.
    for row in rows.values():
        for cell_text in row:
            assert "UNKNOWNPASS" not in cell_text and "UNKNOWNFAIL" not in cell_text


def test_render_candidate_docx_leaves_missing_fields_as_unknown_not_guessed():
    """근거 없는 값을 지어내지 않는다 — 최소 정보만 있는 후보도 안전하게 렌더링된다."""
    candidate = CandidateEquipment(candidate_id="cand-1", source_document="SPEC-999.md")
    docx_bytes = render_candidate_docx(candidate, requirement=None, hard_requirement_report=None)
    document = _open_docx(docx_bytes)

    general_table = document.tables[0]
    rows = {row.cells[0].text: [c.text for c in row.cells] for row in general_table.rows[1:]}
    assert rows["Manufacturer"][1] == "UNKNOWN"
    assert rows["Manufacturer"][2] == "UNKNOWN"


def test_build_candidate_docx_route_returns_file_and_downloads():
    client = _client()
    candidate = _full_candidate()
    resp = client.post(
        "/api/agent/build-candidate-docx",
        json={
            "candidate": candidate.model_dump(),
            "requirement": {"inspection_items": ["thickness"]},
            "hard_requirement_report": [r.model_dump() for r in _hard_requirement_report()],
        },
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["status"] == "success"
    assert data["file_name"].endswith(".docx")
    assert data["file_name"] == "MultiSense_MS-600_specification.docx"
    assert data["download_url"] == f"/api/download/{data['file_name']}"

    download = client.get(data["download_url"])
    assert download.status_code == 200
    assert download.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(download.content) > 0
    document = _open_docx(download.content)
    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "MultiSense MS-600" in full_text


def test_build_candidate_docx_route_works_without_requirement_or_hard_requirement_report():
    client = _client()
    candidate = CandidateEquipment(candidate_id="cand-1", manufacturer="X", model="Y", source_document="SPEC-001.md")
    resp = client.post("/api/agent/build-candidate-docx", json={"candidate": candidate.model_dump()})
    assert resp.status_code == 200


def test_docx_filename_falls_back_when_equipment_has_no_name():
    client = _client()
    candidate = CandidateEquipment(candidate_id="cand-1", source_document="SPEC-999.md")
    resp = client.post("/api/agent/build-candidate-docx", json={"candidate": candidate.model_dump()})
    assert resp.status_code == 200
    assert resp.json()["file_name"] == "equipment_specification.docx"


def test_docx_filename_sanitizes_unsafe_windows_characters():
    client = _client()
    candidate = CandidateEquipment(
        candidate_id="cand-1", manufacturer='Weird/Manu:facturer*Name?"<>|', model="M-1", source_document="SPEC-001.md",
    )
    resp = client.post("/api/agent/build-candidate-docx", json={"candidate": candidate.model_dump()})
    assert resp.status_code == 200
    file_name = resp.json()["file_name"]
    for unsafe in ("/", "\\", ":", "*", "?", '"', "<", ">", "|"):
        assert unsafe not in file_name, f"파일명에 안전하지 않은 문자 '{unsafe}'가 남아있음: {file_name}"
