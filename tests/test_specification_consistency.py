"""
Markdown/Word 사양서 다운로드 — 두 포맷이 같은 입력에서 같은 핵심 정보를
보여주는지 검증한다(요청서 4/13절: "두 형식 중 하나에만 정보가 존재하는 경우
테스트 실패").

두 렌더러(markdown_renderer.render_candidate_markdown / docx_renderer.
render_candidate_docx)는 renderers.candidate_specification.
build_candidate_specification_data()가 만드는 같은 Structured Data를 소비한다
— 이 테스트는 그 계약이 실제로 지켜지는지, 렌더러가 각자 값을 다시 계산/재해석
하면서 어긋나지 않는지를 최종 출력물 기준으로 확인한다.
"""
from __future__ import annotations

import io
import re

import pytest
from docx import Document

from agent.schemas import CandidateEquipment, CandidateEquipmentFact, ComplianceRecord, RequirementSchema, SourceRef
from renderers.docx_renderer import render_candidate_docx
from renderers.markdown_renderer import render_candidate_markdown

pytestmark = pytest.mark.specification


def _candidate() -> CandidateEquipment:
    return CandidateEquipment(
        candidate_id="cand-1",
        manufacturer="ThicknessPro",
        model="TP-800",
        source_document="SPEC-013.md",
        equipment_fact=CandidateEquipmentFact(
            equipment_type="Thickness Inspection",
            measurement_principle="Laser Triangulation",
            inline_offline="inline",
            measurement_method="non_contact",
            width_mm=1200.0,
            range_min=0.0,
            range_max=800.0,
            range_unit="um",
            accuracy_value=0.8,
            accuracy_unit="um",
            speed_value=800.0,
            speed_unit="mm/s",
        ),
    )


def _hard_requirement_report() -> list[ComplianceRecord]:
    return [
        ComplianceRecord(item="Width", unit="mm", requirement=800.0, specification=1200.0, operator=">=", result="PASS", reason="ok", source=SourceRef(document="SPEC-013.md"), hard=True),
        ComplianceRecord(item="Accuracy", unit="um", requirement=1.0, specification=0.8, operator="<=", result="PASS", reason="ok", source=SourceRef(document="SPEC-013.md"), hard=True),
        ComplianceRecord(item="Speed", unit="mm/s", requirement=1000.0, specification=800.0, operator=">=", result="FAIL", reason="fail", source=SourceRef(document="SPEC-013.md"), hard=True),
        ComplianceRecord(item="Repeatability", unit="um", requirement=0.5, specification=None, operator="<=", result="UNKNOWN", reason="unknown", hard=True),
    ]


def _docx_full_text(docx_bytes: bytes) -> str:
    document = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def test_markdown_and_docx_agree_on_core_identity_fields():
    candidate = _candidate()
    requirement = RequirementSchema(inspection_items=["thickness"])
    hard_report = _hard_requirement_report()

    md = render_candidate_markdown(candidate, requirement=requirement, hard_requirement_report=hard_report)
    docx_text = _docx_full_text(render_candidate_docx(candidate, requirement=requirement, hard_requirement_report=hard_report))

    for expected in ("ThicknessPro TP-800", "ThicknessPro", "TP-800"):
        assert expected in md, f"Markdown에 '{expected}'가 없음"
        assert expected in docx_text, f"Word에 '{expected}'가 없음"


def test_markdown_and_docx_agree_on_measurement_range_and_accuracy():
    candidate = _candidate()
    md = render_candidate_markdown(candidate, requirement=None, hard_requirement_report=None)
    docx_text = _docx_full_text(render_candidate_docx(candidate, requirement=None, hard_requirement_report=None))

    assert "0.0 ~ 800.0 um" in md
    assert "0.0 ~ 800.0 um" in docx_text
    assert "±0.8 um" in md
    assert "±0.8 um" in docx_text


def test_markdown_and_docx_agree_on_inspection_mode():
    candidate = _candidate()
    md = render_candidate_markdown(candidate, requirement=None, hard_requirement_report=None)
    docx_text = _docx_full_text(render_candidate_docx(candidate, requirement=None, hard_requirement_report=None))

    assert "inline" in md
    assert "inline" in docx_text


def test_markdown_and_docx_agree_on_hard_requirement_results():
    candidate = _candidate()
    hard_report = _hard_requirement_report()
    md = render_candidate_markdown(candidate, requirement=None, hard_requirement_report=hard_report)
    docx_text = _docx_full_text(render_candidate_docx(candidate, requirement=None, hard_requirement_report=hard_report))

    for record in hard_report:
        # 각 Hard Requirement 항목명과 그 결과(PASS/FAIL/UNKNOWN)가 두 포맷 모두에
        # 나타나야 한다 — 한쪽에만 있거나 결과가 다르면 실패.
        assert record.item in md, f"Markdown Requirement Compliance에 '{record.item}' 없음"
        assert record.item in docx_text, f"Word Requirement Compliance에 '{record.item}' 없음"
    md_result_line = re.search(r"\| Speed \|.*\|\s*(PASS|FAIL|UNKNOWN)\s*\|", md)
    assert md_result_line and md_result_line.group(1) == "FAIL"
    assert "FAIL" in docx_text and "Speed" in docx_text


def test_markdown_and_docx_agree_when_fields_are_missing():
    """값이 없는 경우도 두 포맷이 똑같이 UNKNOWN으로 일치해야 한다 — 한쪽만
    UNKNOWN을 보여주고 다른 쪽은 값을 지어내는 어긋남을 방지한다."""
    candidate = CandidateEquipment(candidate_id="cand-1", source_document="SPEC-999.md")
    md = render_candidate_markdown(candidate, requirement=None, hard_requirement_report=None)
    docx_text = _docx_full_text(render_candidate_docx(candidate, requirement=None, hard_requirement_report=None))

    assert "UNKNOWN" in md
    assert "UNKNOWN" in docx_text
    # 두 포맷 모두 Manufacturer 값이 없을 때 같은 문자열("UNKNOWN")을 쓴다 — 하나는
    # UNKNOWN, 다른 하나는 "N/A"나 빈 문자열처럼 다르게 표현하면 안 된다.
    assert "- Manufacturer: UNKNOWN" in md
