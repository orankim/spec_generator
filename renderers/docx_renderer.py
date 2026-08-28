"""
CandidateEquipment -> Microsoft Word(.docx) 렌더러.

renderers.candidate_specification.build_candidate_specification_data()가 만드는
공통 Structured Data(Markdown 렌더러가 쓰는 것과 동일)를 그대로 소비한다 —
Word와 Markdown이 candidate/requirement를 각자 독립적으로 재해석하지 않도록
하기 위함(요청서 4/11절). 값이 원본 사양서에 없으면 "UNKNOWN"으로 정직하게
남기고, Word 생성 과정에서 추측해서 채우지 않는다(요청서 9절).
"""
from __future__ import annotations

import io
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from agent.schemas import CandidateEquipment, ComplianceRecord, RequirementSchema

from .candidate_specification import CandidateSpecificationData, build_candidate_specification_data

# Hard Requirement 결과 배지 색상 — main.py의 RESULT_BADGE 팔레트(design token)와
# 맞춘다. 새 색상을 만들지 않고 이미 프로젝트가 쓰는 값을 그대로 가져왔다.
_RESULT_COLORS = {
    "PASS": RGBColor(0x1C, 0x6E, 0x7D),
    "FAIL": RGBColor(0x82, 0x27, 0x27),
    "UNKNOWN": RGBColor(0x7B, 0x34, 0x1E),
    "N/A": RGBColor(0x55, 0x55, 0x55),
}


def _style_table(table) -> None:
    table.style = "Light Grid Accent 1"
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True


def _add_section_table(document: Document, section) -> None:
    document.add_heading(section.title, level=1)
    table = document.add_table(rows=1, cols=3)
    _style_table(table)
    header = table.rows[0].cells
    header[0].text, header[1].text, header[2].text = "Item", "Specification", "Status"
    for row in section.rows:
        cells = table.add_row().cells
        cells[0].text = row.label
        cells[1].text = row.value
        cells[2].text = row.status
        if row.status == "UNKNOWN":
            for cell in (cells[1], cells[2]):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = _RESULT_COLORS["UNKNOWN"]
    document.add_paragraph()


def _add_compliance_table(document: Document, compliance) -> None:
    document.add_heading("Requirement Compliance", level=1)
    if not compliance:
        document.add_paragraph("No requirement provided for comparison.")
        return
    table = document.add_table(rows=1, cols=4)
    _style_table(table)
    header = table.rows[0].cells
    header[0].text, header[1].text, header[2].text, header[3].text = "Requirement", "Required", "Equipment", "Result"
    for row in compliance:
        cells = table.add_row().cells
        cells[0].text = row.item
        cells[1].text = row.required_display
        cells[2].text = row.equipment_display
        cells[3].text = row.result
        color = _RESULT_COLORS.get(row.result)
        if color is not None:
            for paragraph in cells[3].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = color
    document.add_paragraph()


def _build_document(data: CandidateSpecificationData) -> Document:
    document = Document()

    title = document.add_heading(data.equipment_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # General Specification은 표로 별도 렌더링(다른 섹션과 동일한 형식으로 통일).
    general = next(s for s in data.sections if s.id == "general")
    _add_section_table(document, general)

    for section in data.sections:
        if section.id == "general":
            continue
        _add_section_table(document, section)

    _add_compliance_table(document, data.compliance)

    document.add_heading("Sources / Notes", level=1)
    if data.sources:
        document.add_paragraph("Reference Documents:")
        for source in data.sources:
            document.add_paragraph(source, style="List Bullet")
    else:
        document.add_paragraph("Reference Documents: UNKNOWN")
    if data.notes:
        document.add_paragraph("Notes:")
        for note in data.notes:
            document.add_paragraph(note, style="List Bullet")

    return document


def render_candidate_docx(
    candidate: CandidateEquipment,
    requirement: Optional[RequirementSchema] = None,
    hard_requirement_report: Optional[List[ComplianceRecord]] = None,
) -> bytes:
    """CandidateEquipment 하나를 Word(.docx) 바이트로 렌더링한다. render_candidate_
    markdown()과 정확히 같은 build_candidate_specification_data() 결과를 쓰므로
    두 포맷의 핵심 정보(장비명/Manufacturer/Model/Range/Accuracy/Inspection Mode/
    Hard Requirement 결과)가 항상 일치한다."""
    data = build_candidate_specification_data(candidate, requirement=requirement, hard_requirement_report=hard_requirement_report)
    document = _build_document(data)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
